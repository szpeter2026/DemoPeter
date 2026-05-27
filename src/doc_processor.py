"""
szpeter2026 - 文档处理管道
吸收自 HKIE/Wukong 的文档处理模式：PDF/MD/TXT/DOCX/IPYNB/CSV → 文本提取 → 智能分块
"""
import csv
import json
import re
from pathlib import Path
from typing import Iterator

from config.settings import config


class DocumentProcessor:
    """文档处理管道 — 导入 → 提取 → 分块"""

    SUPPORTED_TYPES = {
        ".md", ".markdown", ".txt", ".pdf",
        ".docx", ".ipynb", ".csv",
    }
    SKIP_SUFFIXES = {".zip", ".xlsx", ".xls", ".pptx"}
    # 跳过嵌套解压副本目录（数据分析实战代码包内重复一层）
    SKIP_DIR_NAMES = {"__pycache__", ".git", ".ipynb_checkpoints"}

    def __init__(self):
        self.chunk_size = config.CHUNK_SIZE
        self.chunk_overlap = config.CHUNK_OVERLAP

    # ===== 文本提取 =====

    @staticmethod
    def extract_text(file_path: str | Path) -> str:
        """从文件提取文本内容"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix in (".md", ".markdown"):
            return file_path.read_text(encoding="utf-8")
        if suffix == ".txt":
            return DocumentProcessor._read_text_with_encodings(file_path)
        if suffix == ".pdf":
            return DocumentProcessor._extract_pdf(file_path)
        if suffix == ".docx":
            return DocumentProcessor._extract_docx(file_path)
        if suffix == ".ipynb":
            return DocumentProcessor._extract_ipynb(file_path)
        if suffix == ".csv":
            return DocumentProcessor._extract_csv(file_path)
        raise ValueError(f"不支持的文件类型: {suffix}")

    @staticmethod
    def _read_text_with_encodings(file_path: Path) -> str:
        for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
            try:
                return file_path.read_text(encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return file_path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        """PDF 文本提取 — 优先使用 pdfplumber，回退 pypdf"""
        text_parts = []
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            if text_parts:
                return "\n\n".join(text_parts)
        except ImportError:
            pass

        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("处理 .docx 需要安装 python-docx") from exc
        doc = Document(str(file_path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    @staticmethod
    def _extract_ipynb(file_path: Path) -> str:
        """提取 Jupyter Notebook 中 markdown + code 单元格"""
        raw = json.loads(file_path.read_text(encoding="utf-8"))
        parts = []
        title = raw.get("metadata", {}).get("kernelspec", {}).get("display_name", file_path.stem)
        parts.append(f"# Notebook: {title}")
        for i, cell in enumerate(raw.get("cells", []), 1):
            src = cell.get("source", "")
            if isinstance(src, list):
                src = "".join(src)
            src = src.strip()
            if not src:
                continue
            cell_type = cell.get("cell_type", "code")
            parts.append(f"## Cell {i} ({cell_type})\n{src}")
        return "\n\n".join(parts)

    @staticmethod
    def _extract_csv(file_path: Path, max_rows: int = 200) -> str:
        """CSV → 可读文本（表头 + 采样行）"""
        lines = [f"# CSV: {file_path.name}"]
        with file_path.open(encoding="utf-8", errors="replace", newline="") as fh:
            reader = csv.reader(fh)
            rows = list(reader)
        if not rows:
            return lines[0]
        header = rows[0]
        lines.append("列: " + ", ".join(header))
        for row in rows[1:max_rows + 1]:
            pairs = []
            for h, v in zip(header, row):
                pairs.append(f"{h}={v}")
            lines.append(" | ".join(pairs))
        if len(rows) > max_rows + 1:
            lines.append(f"... 共 {len(rows) - 1} 行，仅展示前 {max_rows} 行")
        return "\n".join(lines)

    # ===== 智能分块 =====

    def chunk_text(self, text: str, metadata: dict | None = None) -> list[dict]:
        """将文本智能分块，保留段落/章节边界"""
        meta = metadata or {}
        chunks = []
        paragraphs = self._split_by_structure(text)

        current_chunk = ""
        for para in paragraphs:
            if not para.strip():
                continue

            if len(current_chunk) + len(para) <= self.chunk_size:
                current_chunk += para + "\n\n"
            else:
                # 保存当前块
                if current_chunk.strip():
                    chunks.append({
                        "index": len(chunks),
                        "content": current_chunk.strip(),
                        "metadata": {**meta, "char_count": len(current_chunk.strip())}
                    })
                # 开始新块（带重叠）
                overlap_text = current_chunk[-self.chunk_overlap:] if self.chunk_overlap and current_chunk else ""
                current_chunk = overlap_text + para + "\n\n"

        # 最后一个块
        if current_chunk.strip():
            chunks.append({
                "index": len(chunks),
                "content": current_chunk.strip(),
                "metadata": {**meta, "char_count": len(current_chunk.strip())}
            })

        return chunks

    @staticmethod
    def _split_by_structure(text: str) -> list[str]:
        """按 Markdown 标题和自然段落分割"""
        # 先按 Markdown 标题分割
        sections = re.split(r'(?=^#{1,6}\s)', text, flags=re.MULTILINE)
        result = []
        for section in sections:
            # 再在内部按双换行分割
            paras = re.split(r'\n\s*\n', section)
            result.extend(paras)
        return result

    # ===== 批量导入 =====

    def scan_directory(
        self,
        directory: str | Path,
        *,
        corpus_label: str = "",
        skip_nested_duplicate: bool = True,
    ) -> list[dict]:
        """扫描目录下所有可处理的文件"""
        directory = Path(directory).resolve()
        root_name = directory.name
        files = []
        seen_paths: set[str] = set()

        for file_path in sorted(directory.rglob("*")):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() in self.SKIP_SUFFIXES:
                continue
            if file_path.suffix.lower() not in self.SUPPORTED_TYPES:
                continue
            if any(part in self.SKIP_DIR_NAMES for part in file_path.parts):
                continue
            # 跳过「数据分析与python实战-代码/数据分析与python实战-代码/」嵌套副本
            if skip_nested_duplicate and root_name in file_path.parts:
                parts = file_path.parts
                if parts.count(root_name) > 1:
                    continue

            resolved = str(file_path.resolve())
            if resolved in seen_paths:
                continue
            seen_paths.add(resolved)

            files.append({
                "title": file_path.stem,
                "file_path": resolved,
                "doc_type": file_path.suffix.lower().lstrip("."),
                "file_size": file_path.stat().st_size,
                "corpus_label": corpus_label or root_name,
            })
        return files

    def process_file(
        self,
        file_path: str | Path,
        extra_metadata: dict | None = None,
    ) -> tuple[str, list[dict]]:
        """处理单个文件：提取文本 + 分块"""
        text = self.extract_text(file_path)
        file_path = Path(file_path)
        meta = {
            "source_file": file_path.name,
            "source_path": str(file_path.parent),
            "doc_type": file_path.suffix.lower().lstrip("."),
            **(extra_metadata or {}),
        }
        chunks = self.chunk_text(text, metadata=meta)
        return text, chunks
